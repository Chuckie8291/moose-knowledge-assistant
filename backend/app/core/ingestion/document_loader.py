"""
Document Loader — Loads raw files from disk/S3 and extracts raw text.

Supports: PDF (digital + scanned), DOCX, TXT, PNG/JPG/TIFF.
"""

from __future__ import annotations

import hashlib
import io
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from app.utils.logging import get_logger

logger = get_logger(__name__)


@dataclass
class LoadedPage:
    """A single page of extracted content."""
    page_number: int          # 1-indexed physical page
    logical_page_number: Optional[str]  # e.g., "iii", "47" (from page label)
    text: str                 # Extracted or OCR'd text
    char_count: int
    image_bytes: Optional[bytes]  # Raw page image (for OCR pipeline)


@dataclass
class LoadedDocument:
    """Result of loading a document from file."""
    file_path: str
    file_ext: str
    file_size_bytes: int
    content_hash: str          # SHA-256 of raw file
    pages: list[LoadedPage]
    total_pages: int
    total_chars: int
    is_digital: bool           # True = text extractable; False = needs OCR
    metadata: dict             # Any file-level metadata (PDF info, DOCX properties)


class DocumentLoader:
    """Load and extract text from various document formats."""

    MAX_PREVIEW_PAGES = 5     # For classification preview

    def load(self, file_path: str) -> LoadedDocument:
        """Main entry point. Auto-detects format and delegates."""
        path = Path(file_path)
        ext = path.suffix.lower()
        file_bytes = path.read_bytes()
        content_hash = hashlib.sha256(file_bytes).hexdigest()

        if ext == '.pdf':
            return self._load_pdf(file_path, file_bytes, content_hash)
        elif ext in ('.docx', '.doc'):
            return self._load_docx(file_path, file_bytes, content_hash)
        elif ext in ('.txt', '.text'):
            return self._load_text(file_path, file_bytes, content_hash)
        elif ext in ('.png', '.jpg', '.jpeg', '.tiff', '.tif'):
            return self._load_image(file_path, file_bytes, content_hash)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    # ── PDF ──────────────────────────────────────────────────

    def _load_pdf(
        self, file_path: str, file_bytes: bytes, content_hash: str
    ) -> LoadedDocument:
        """Load PDF — try text extraction first, flag for OCR if needed."""
        import fitz  # pymupdf

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        total_chars = 0
        digital_page_count = 0

        for i, page in enumerate(doc):
            text = page.get_text("text") or ""
            char_count = len(text.strip())

            # Extract logical page number from PDF labels
            logical = None
            label = page.get_label()
            if label and label != str(i + 1):
                logical = label

            # Render page as image (for potential OCR)
            pix = page.get_pixmap(dpi=200)
            image_bytes = pix.tobytes("png")

            pages.append(LoadedPage(
                page_number=i + 1,
                logical_page_number=logical,
                text=text.strip(),
                char_count=char_count,
                image_bytes=image_bytes,
            ))
            total_chars += char_count
            if char_count >= 50:
                digital_page_count += 1

        doc.close()

        is_digital = (digital_page_count / max(len(pages), 1)) >= 0.90

        metadata = {
            "format": "PDF",
            "pdf_page_count": len(pages),
            "pdf_title": doc.metadata.get("title", ""),
            "pdf_author": doc.metadata.get("author", ""),
        }

        return LoadedDocument(
            file_path=file_path,
            file_ext=".pdf",
            file_size_bytes=len(file_bytes),
            content_hash=content_hash,
            pages=pages,
            total_pages=len(pages),
            total_chars=total_chars,
            is_digital=is_digital,
            metadata=metadata,
        )

    # ── DOCX ─────────────────────────────────────────────────

    def _load_docx(
        self, file_path: str, file_bytes: bytes, content_hash: str
    ) -> LoadedDocument:
        """Load Word document — extract text and heading structure."""
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))

        # Extract all paragraphs with style info
        all_text_parts = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            is_heading = "Heading" in style_name or "heading" in style_name
            all_text_parts.append({
                "text": para.text,
                "is_heading": is_heading,
                "style": style_name,
                "heading_level": self._parse_heading_level(style_name),
            })

        full_text = "\n".join(p["text"] for p in all_text_parts)
        total_chars = len(full_text.strip())

        # DOCX doesn't have pages — treat as single logical page
        page = LoadedPage(
            page_number=1,
            logical_page_number="1",
            text=full_text.strip(),
            char_count=total_chars,
            image_bytes=None,
        )

        return LoadedDocument(
            file_path=file_path,
            file_ext=".docx",
            file_size_bytes=len(file_bytes),
            content_hash=content_hash,
            pages=[page],
            total_pages=1,
            total_chars=total_chars,
            is_digital=True,
            metadata={
                "format": "DOCX",
                "paragraphs": len(all_text_parts),
                "heading_structure": [
                    p for p in all_text_parts if p["is_heading"]
                ],
            },
        )

    @staticmethod
    def _parse_heading_level(style_name: str) -> int:
        """Extract heading level from Word style name. 'Heading 2' → 2."""
        import re
        match = re.search(r'(\d+)', style_name)
        return int(match.group(1)) if match else 0

    # ── Plain Text ───────────────────────────────────────────

    def _load_text(
        self, file_path: str, file_bytes: bytes, content_hash: str
    ) -> LoadedDocument:
        """Load plain text file."""
        text = file_bytes.decode("utf-8", errors="replace")
        page = LoadedPage(
            page_number=1,
            logical_page_number="1",
            text=text.strip(),
            char_count=len(text.strip()),
            image_bytes=None,
        )
        return LoadedDocument(
            file_path=file_path,
            file_ext=".txt",
            file_size_bytes=len(file_bytes),
            content_hash=content_hash,
            pages=[page],
            total_pages=1,
            total_chars=len(text.strip()),
            is_digital=True,
            metadata={"format": "TXT"},
        )

    # ── Image ────────────────────────────────────────────────

    def _load_image(
        self, file_path: str, file_bytes: bytes, content_hash: str
    ) -> LoadedDocument:
        """Load image — always needs OCR."""
        # For single images, create a page with empty text (OCR will fill it in)
        page = LoadedPage(
            page_number=1,
            logical_page_number=None,
            text="",  # Will be populated by OCR
            char_count=0,
            image_bytes=file_bytes,
        )
        return LoadedDocument(
            file_path=file_path,
            file_ext=Path(file_path).suffix.lower(),
            file_size_bytes=len(file_bytes),
            content_hash=content_hash,
            pages=[page],
            total_pages=1,
            total_chars=0,
            is_digital=False,  # Images always need OCR
            metadata={"format": "IMAGE"},
        )

    # ── Preview (for classification) ─────────────────────────

    def load_preview(self, file_path: str) -> str:
        """Extract first few pages of text for document classification."""
        loaded = self.load(file_path)
        preview_pages = loaded.pages[:self.MAX_PREVIEW_PAGES]
        return "\n\n".join(p.text for p in preview_pages if p.text)
